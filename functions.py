import os
import docx
import json
from parse import Parser


def get_paths(folder):
    """Возрващает список путей к файлам .docx в заданной директории."""

    # folder = os.getcwd()
    paths = []
    for root, dirs, files in os.walk(folder):
        for file in files:
            if file.endswith('docx') and not file.startswith('~') and not file.startswith('default'):
                paths.append(os.path.join(root, file))
    return paths


def get_paragraphs_multiple(paths):
    """Возвращает список документов, элементами которого являются параграфы каждого документа"""

    paragraphs = []
    for path in paths:
        doc = docx.Document(path)
        paragraphs.append(doc.paragraphs)
    return paragraphs


def get_paragraphs_single(path):
    """Возвращает список параграфов одного документа"""

    doc = docx.Document(path)
    return doc.paragraphs


def get_pattern(filename):
    """Чтение шаблона в формате .json"""

    with open(filename, encoding='utf-8') as json_file:
        data = json.load(json_file)
    return data


def parse_doc(path_to_parse, path_to_save):
    """Парсинг одного документа по заданному пути"""

    doc = get_paragraphs_single(path_to_parse)
    parser = Parser(doc)
    parsed_data = parser.parse()
    outfile = path_to_parse.split('\\')[-1].split('.')[0] + '.json'
    # parser.to_json_file(f"./ParsedResults/{outfile}")
    parser.to_json_file(path_to_save + "\\" + outfile)
    return parsed_data


def parse_docs():
    """Парсинг всех .docx документов в корневой папке проекта"""

    paths = get_paths()
    docs = get_paragraphs_multiple(paths)
    data = []
    for doc, path in zip(docs, paths):
        parser = Parser(doc)
        parsed_data = parser.parse()
        data.append(parsed_data)
        print(json.dumps(Parser.to_json(parsed_data), indent=4, ensure_ascii=False))
        outfile = path.split('/')[-1].split(".")[0] + ".json"
        parser.to_json_file(f"./ParsedResults//{outfile}")
    return data


def test_to_json(data, path):
    with open(path, 'w', encoding='utf-8') as outfile:
        json.dump(data, outfile, indent=4, ensure_ascii=False)


def calc_result(headings_check, bodies_check):
    """Результат всех проверок -- линейная комбинация всех показателей"""

    missing_partitions = headings_check['missing partitions']
    unordered_number = headings_check['number of unordered partitions']
    headings_avg = sum(headings_check.values()) / len(headings_check)
    bodies_avg = sum(bodies_check.values()) / len(bodies_check)

    values_list = [
        missing_partitions * 0.5,
        unordered_number * 0.5,
        headings_avg * 0.5,
        bodies_avg * 0.4
    ]
    return sum(values_list)
